import io,json
import pathlib
from docx import Document
from docx.shared import Inches,Mm,Pt
import datetime
from pathlib import Path
import pandas as pd
from pandas import DataFrame
import math
import numpy as np
import plotly.graph_objects as go

import honeybee_ies as hb2ies
from honeybee.model import Model as HBModel
from ladybug_geometry.geometry2d.pointvector import Vector2D
from honeybee_vtk.model import Model as VTKModel

from pollination_streamlit_viewer import viewer
from pollination_streamlit_io import get_hbjson


import streamlit as st

st.set_page_config(page_title='SpaceXtract',
    layout="wide"
)

hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                </style>
                """

#Side BAR
with st.sidebar:
    st.image('./img/SpaceXtract.png',use_column_width='auto',output_format='PNG',
             caption='This tool accepts .json files exported from Ladybyug-tools and will automatically extracts building envelope information. You can also export the 3D model into a .gem file for IES users. More features are to be developed!')

    st.header("User Control Panel:")

    north_ = st.number_input("**North Angle (90° by default):**", -180,180, 90 , 1, key = 'north',help = "Counter-Clockwise for Negative Values, Clockwise for Positive Values")
    vectors = [math.cos(math.radians(north_)), math.sin(math.radians(north_))]
    

#MAIN PAGE
st.title("SpaceXtract")
st.subheader("Before uploading, please check the following items:")
st.markdown("**1- Each room should have unique names to calculate the internal wall surface areas accurately!**")
st.markdown("**2- Boundary conditions of adjacent surfaces should be done prior to use the tool, otherwise external wall surface areas will include internal walls!**")
st.markdown("**3- If you are interested to calculate the conditioned spaces only, they should be set as conditioned in the model!**")

st.subheader("Upload .hbjson Model:")

if 'temp' not in st.session_state:
     st.session_state.temp = Path('/tmp')
     st.session_state.temp.mkdir(parents=True, exist_ok=True)

def create_vtkjs(hbjson_path: Path):
    if not hbjson_path:
        return
    
    model = VTKModel.from_hbjson(hbjson_path.as_posix())
    
    vtkjs_dir = st.session_state.temp.joinpath('vtkjs')
    if not vtkjs_dir.exists():
        vtkjs_dir.mkdir(parents=True, exist_ok=True)

    vtkjs_file = vtkjs_dir.joinpath(f'{hbjson_path.stem}.vtkjs')
    
    if not vtkjs_file.is_file():
        model.to_vtkjs(
            folder=vtkjs_dir.as_posix(),
            name=hbjson_path.stem
        )

    return vtkjs_file

def show_model(hbjson_path: Path):
    """Render HBJSON."""

    vtkjs_name = f'{hbjson_path.stem}_vtkjs'
    vtkjs = create_vtkjs(hbjson_path)
    st.session_state.content = vtkjs.read_bytes()
    st.session_state[vtkjs_name] = vtkjs
    
    
def callback_once():
    if 'hbjson' in st.session_state.get_hbjson:
        hb_model = HBModel.from_dict(st.session_state.get_hbjson['hbjson'])
        if hb_model:
            hbjson_path = st.session_state.temp.joinpath(f'{hb_model.identifier}.hbjson')
            hbjson_path.write_text(json.dumps(hb_model.to_dict()))
            show_model(hbjson_path)
    return hb_model


hbjson = get_hbjson('get_hbjson', on_change=callback_once)

if st.session_state.get_hbjson is not None:
    st.subheader(f'Visualizing {callback_once().display_name} Model')
    if 'content' in st.session_state:
        viewer(
            content=st.session_state.content,
            key='vtkjs-viewer',
            subscribe=False,
            style={
                'height' : '640px'
            }
        )

    exports = st.columns([5,5,10])

    with exports[0]:
        # EXPORT AS GEM FILE
        path_to_ies_folder = pathlib.Path('./gem')
        path_to_ies_folder.mkdir(parents=True, exist_ok=True)

        if st.button("**Export as .GEM File (for IES Users)**"):
                
            data=hb2ies.writer.model_to_ies(model =callback_once(), folder = path_to_ies_folder)

    with exports[1]:
        #EXPORT AS IDF FILE
        if st.button("**Export as .IDF File (for ENERGYPLUS Users)**"):
                
            data=callback_once().to.idf(callback_once())
            with open(f'./idf/{callback_once().display_name}.idf', 'w') as file:
                file.write(data)
else:
    st.info('Load a model!')

#Generating the model

with st.sidebar:

    area_calc_method = st.radio("**Select the Facade Area Calculation Methodology**", options = ['Conditioned Zones', 'Entire Building'])

    baseline_calc = st.radio("**Select the Baseline Energy Calculation Methodology**", options = ['None','Facade Calculator NCC 2019 (Australia)', 'ASHRAE 90.1 Guidelines', 'PassivHause Standard'])

    bldg_classes = {'Class 2 - apartment building (Common Area)':2,
    'Class 3 - student accommodation':3,
    'Class 3 - hotel':3,
    'Class 3 - other':3,
    'Class 5 - office building':5,
    'Class 6 - department stores, shopping centres':6,
    'Class 6 - display glazing':6,
    'Class 6 - restaurants, cafes, bars':6,
    'Class 8 - factory':8,
    'Class 9a - health-care buildings':'9a',
    'Class 9a - ward':'9a ward',
    'Class 9b - churches, chapels or the like':'9b',
    'Class 9b - early childhood centres':'9b',
    'Class 9b - public halls, function rooms or the like':'9b',
    'Class 9b - schools':'9b',
    'Class 9b - single auditorium theatres and cinemas':'9b',
    'Class 9b - sports venues or the like':'9b',
    'Class 9b - theatres and cinemas with multiple auditoria, art galleries or the like':'9b',
    'Class 9c - aged care building':'9c'}

    aus_climate_zone = {'Climate Zone 1 - High humidity summer, warm winter':1,
    'Climate Zone 2 - Warm humid summer, mild winter':2,
    'Climate Zone 3 - Hot dry summer, warm winter':3,
    'Climate Zone 4 - Hot dry summer, cool winter':4,
    'Climate Zone 5 - Warm temperate':5,
    'Climate Zone 6 - Mild temperate':6,
    'Climate Zone 7 - Cool temperate':7,
    'Climate Zone 8 - Alpine':8}

    if baseline_calc == 'None':
        ""

    elif baseline_calc == 'Facade Calculator NCC 2019 (Australia)':
        building_state = st.selectbox('**Building State:**',['ACT','NT','QLD','NSW','SA','TAS','VIC','WA'], index = 2)
        building_class = st.selectbox('**Building Classification:**',bldg_classes, index = 4)
        climate_zone = st.selectbox('**Climate Zone:**',aus_climate_zone, index = 1)
        ex_wall_dts = st.number_input('**External Wall R-value:**', value = 1.4)
        glass_u_dts = st.number_input('**Glass U-value:**', value = 3.5)
        glass_shgc_dts = st.number_input('**Glass SHGC:**', value = 0.5)

    elif baseline_calc == 'ASHRAE Standard (US)':
        ""


if st.session_state.get_hbjson is not None:
    model = callback_once()

    #Extracting properties based on rooms
    model_data = {'Conditioning Status':[], 'Program Type': [], 'volume (m3)': [],'floor_area (m2)': [],'roof_area (m2)':[],  'exterior_wall_area (m2)': [],
                    'exterior_aperture_area (m2)': [], 'exterior_skylight_area (m2)':[],'display_name':[]}

    internal_srf_area = []
    room_id = []
    for room in model.rooms:
        model_data['display_name'].append(room.display_name)
        model_data['Conditioning Status'].append(room.properties.energy.is_conditioned)
        model_data['Program Type'].append(room.properties.energy.program_type._identifier)
        model_data['volume (m3)'].append(room.volume)
        model_data['floor_area (m2)'].append(room.floor_area)
        model_data['roof_area (m2)'].append(room.exterior_roof_area)
        model_data['exterior_wall_area (m2)'].append(room.exterior_wall_area - room.exterior_aperture_area)
        model_data['exterior_aperture_area (m2)'].append(room.exterior_aperture_area)
        model_data['exterior_skylight_area (m2)'].append(room.exterior_skylight_aperture_area)
        #internal walls   
        for face in room.faces:
            room_id.append(room.display_name)
            if face.boundary_condition.name == 'Surface' and face.azimuth > 0:
                internal_srf_area.append(face.area)
            else:
                internal_srf_area.append(0)
        internal = DataFrame([room_id,internal_srf_area],['ROOM_ID','Internal_Faces']).transpose().sort_values('ROOM_ID').groupby('ROOM_ID').sum()

    model_data = DataFrame.from_dict(model_data).sort_values('display_name').set_index('display_name')

    model_data['internal_wall_area (m2)'] = internal['Internal_Faces']

    model_shade = {'total_external_shades_area (m2)':[]}

    for shade in model.outdoor_shades:
        model_shade['total_external_shades_area (m2)'].append(shade.area)

    model_shade = DataFrame.from_dict(model_shade).sum()

    #Extracting room index based on facade calc methodology

    target_rooms_index = []

    for i,room in enumerate(range(len(model_data.index))):
        
        if (area_calc_method == 'Conditioned Zones') and (model_data['Conditioning Status'].iloc[room] == True):
            target_rooms_index.append(i)
            
        elif area_calc_method == 'Entire Building':
            target_rooms_index.append(i)

    #Surface Areas based on directions
    roof_faces_area = []
    floor_faces_area = []
    vert_face_area = []
    face_orientation = []

    # aperture_orientation = {'North':[],'East':[],'South':[],'West':[]}
    aperture_orientation = {'North':[],'East':[],'South':[],'West':[]}
    aperture_area_north = []
    aperture_area_south = []
    aperture_area_east = []
    aperture_area_west = []

    model_roof = {'Roof Area (m2)':[]}
    model_floor = {'Floor Area (m2)':[]}

    for room_index in target_rooms_index:
        
        for aperture in range(len(model.rooms[room_index].exterior_apertures)):
            if model.rooms[room_index].exterior_apertures[aperture].azimuth > 0: #excluding skylights if any
                
                aper_azimuth = model.rooms[room_index].exterior_apertures[aperture].horizontal_orientation(north_vector=Vector2D(vectors[0],vectors[1]))

                if (aper_azimuth <= 45) or (aper_azimuth > 315):
                    aperture_area_north.append(model.rooms[room_index].exterior_apertures[aperture].area)
                    
                elif (aper_azimuth > 45) and (aper_azimuth <= 135):
                    aperture_area_east.append(model.rooms[room_index].exterior_apertures[aperture].area)
                    
                elif (aper_azimuth > 135) and (aper_azimuth <= 225):
                    aperture_area_south.append(model.rooms[room_index].exterior_apertures[aperture].area)
                    
                elif (aper_azimuth > 225) and (aper_azimuth <= 315):
                    aperture_area_west.append(model.rooms[room_index].exterior_apertures[aperture].area)
                    
        aperture_orientation['North'] = round(sum(aperture_area_north),2)
        aperture_orientation['East'] = round(sum(aperture_area_east),2)
        aperture_orientation['South'] = round(sum(aperture_area_south),2)
        aperture_orientation['West'] = round(sum(aperture_area_west),2)
        model_apertures = DataFrame.from_dict([aperture_orientation]).transpose()
        model_apertures.rename(columns = {0:'Aperture Area (m2)'}, inplace= True)
        model_apertures = model_apertures.sort_index()

        for face in range(len(model.rooms[room_index].faces)):
            if (model.rooms[room_index].faces[face].boundary_condition.name == 'Outdoors') and (model.rooms[room_index].faces[face].type.name != 'RoofCeiling') and (model.rooms[room_index].faces[face].type.name != 'Floor'):
            
                vert_face_azimuth = model.rooms[room_index].faces[face].horizontal_orientation(north_vector=Vector2D(vectors[0],vectors[1]))

                if vert_face_azimuth <= 45 or vert_face_azimuth > 315:
                    face_orientation.append('North')
                    vert_face_area.append(round(model.rooms[room_index].faces[face].area,2))

                elif vert_face_azimuth > 45 and vert_face_azimuth <= 135:
                    face_orientation.append('East')                                
                    vert_face_area.append(round(model.rooms[room_index].faces[face].area,2))

                elif vert_face_azimuth > 135 and vert_face_azimuth <= 225:
                    face_orientation.append('South')
                    vert_face_area.append(round(model.rooms[room_index].faces[face].area,2))
            
                elif vert_face_azimuth > 225 and vert_face_azimuth <= 315:
                    face_orientation.append('West')
                    vert_face_area.append(round(model.rooms[room_index].faces[face].area,2))

            model_faces_vertical = DataFrame([face_orientation,vert_face_area]).transpose()
            model_faces_vertical.rename(columns = {0:'Face Orientation', 1:'Face Area (m2)'}, inplace= True)
            model_faces_vertical = model_faces_vertical.groupby('Face Orientation').sum()
            model_faces_vertical = model_faces_vertical.sort_index()
            

            #checking horizontal faces
            if model.rooms[room_index].faces[face].type.name == 'RoofCeiling':
                horiz_face_area = model.rooms[room_index].faces[face].area
                roof_faces_area.append(round(horiz_face_area,2))
                
            if model.rooms[room_index].faces[face].type.name == 'Floor': #if there is any exposed floor, if not returns 0
                horiz_face_area = model.rooms[room_index].faces[face].area
                floor_faces_area.append(round(horiz_face_area,2))
                
        
        model_faces_vertical['ExWall Area (m2)'] = round(model_faces_vertical['Face Area (m2)'] - model_apertures['Aperture Area (m2)'],2)
        model_faces_vertical['WWR (%)'] = round((model_apertures['Aperture Area (m2)'] / model_faces_vertical['Face Area (m2)'])*100,2)
        model_faces_vertical['WWR (%)'].fillna(0, inplace=True)
        model_roof['Roof Area (m2)'] = roof_faces_area    
        model_roof_DF= DataFrame.from_dict(model_roof).sum()
        model_floor['Floor Area (m2)'] = floor_faces_area
        model_floor_DF = DataFrame.from_dict(model_floor).sum()

 

if st.session_state.get_hbjson is not None:
    #Building Relative Compactness (RC) = 6 * Building Volume (V) ^ 2/3 / Building Surface Area (A)
    ##Source: https://www.sciencedirect.com/science/article/abs/pii/S037877881400574X?via%3Dihub

    building_volume = model_data['volume (m3)'].sum()
    Building_area = model_data[['floor_area (m2)','roof_area (m2)','exterior_wall_area (m2)','exterior_aperture_area (m2)','exterior_skylight_area (m2)']].sum().sum() #internal walls excluded 

    build_RC = (6 * (pow(building_volume,2/3))) / Building_area

    st.header(f'**Building Relative Compactness (RC)** is :red[{round(build_RC,2)}].')

    st.markdown('---')

    st.subheader(f'**Building General Details**')
    
    if model_data.index.nunique() != len(model_data.index): #Checking room names similarity for internal walls calculations
        st.warning("There are similar room names in the model which will cause in inaccurate internal wall surface areas calculations! Please fix them before uploading the model.")
    else:
        ""
    st.dataframe(model_data, use_container_width=True)
    st.markdown('---')
    
    st.subheader(f'**Thermal Envelope Area Calculation Based on {area_calc_method}**')
    
    if target_rooms_index == []:
        st.subheader(":red[OOPS! NO CONDITIONED ZONES ARE ASSIGENED IN THE MODEL!]")
    else:
        cols = st.columns(4)

        with cols[0]:
            st.dataframe(model_shade, use_container_width=True)
            
            col_wwr = st.columns(len(model_faces_vertical.index))

            for metric in range(len(model_faces_vertical.index)):
            
                with col_wwr[metric]:
                    st.metric(f"WWR-{model_faces_vertical.index[metric]}",f"{int(model_faces_vertical['WWR (%)'].iloc[metric])}%")
        
        with cols[1]:
            st.dataframe(model_apertures, use_container_width=True)
        with cols[2]:
            st.dataframe(model_faces_vertical.drop(['Face Area (m2)','WWR (%)'], axis = 1), use_container_width=True)
        with cols[3]:
            st.dataframe(model_roof_DF, use_container_width=True)
            st.dataframe(model_floor_DF, use_container_width=True)

    st.markdown('---')


#DtS Facade Calculation NCC2019 (AUSTRALIA)
if st.session_state.get_hbjson is not None and target_rooms_index != []:
    if baseline_calc == 'Facade Calculator NCC 2019 (Australia)':
        st.header("Reference Building Fabric Performance - NCC19 Facade Calculator")
        
        #R Target
        R_target = []
        for i in range(0,4):
            if model_faces_vertical['WWR (%)'].iloc[i] <= 20:
                if bldg_classes[building_class] == 2 or bldg_classes[building_class] == 5 or bldg_classes[building_class] == 6 or bldg_classes[building_class] == 7 or bldg_classes[building_class] == 8 or bldg_classes[building_class] == '9b' or bldg_classes[building_class] == '9a':
                    if aus_climate_zone[climate_zone] == 2 or aus_climate_zone[climate_zone] == 3 or aus_climate_zone[climate_zone] == 4 or aus_climate_zone[climate_zone] == 5 or aus_climate_zone[climate_zone] == 6 or aus_climate_zone[climate_zone] == 7 or aus_climate_zone[climate_zone] == 8:
                        R_target.append(1.4)
                    elif aus_climate_zone[climate_zone] == 1:
                        R_target.append(2.4)
                elif bldg_classes[building_class] == 3 or bldg_classes[building_class] == '9c' or bldg_classes[building_class] == '9a ward':
                    if aus_climate_zone[climate_zone] == 1 or aus_climate_zone[climate_zone] == 3:
                        R_target.append(3.3)
                    elif aus_climate_zone[climate_zone] == 2 or aus_climate_zone[climate_zone] == 5:
                        R_target.append(1.4)
                    elif aus_climate_zone[climate_zone] == 4 or aus_climate_zone[climate_zone] == 6 or aus_climate_zone[climate_zone] == 7:
                        R_target.append(2.8)
                    elif aus_climate_zone[climate_zone] == 8:
                        R_target.append(3.8)
            elif model_faces_vertical['WWR (%)'].iloc[i] > 20:
                R_target.append(1.0)

        Wall_U_Value = []
        for i in range(0,4):
            if model_faces_vertical['WWR (%)'].iloc[i] <= 20 and 1/ex_wall_dts >= 1/R_target[i]:
                Wall_U_Value.append(1/R_target[i]) 
            else:
                if model_faces_vertical['WWR (%)'].iloc[i] > 20 and 1/ex_wall_dts >= 1/R_target[i]:
                    Wall_U_Value.append(1/R_target[i])
                else:
                    Wall_U_Value.append(1/ex_wall_dts)


        shading_multi = 1.0 #assuming no shades for reference buildings and will be equal for all orientations

        #Target Wall Glazing U-values
        if bldg_classes[building_class] == 2 or bldg_classes[building_class] == 5 or bldg_classes[building_class] == 6 or bldg_classes[building_class] == 7 or bldg_classes[building_class] == 8 or bldg_classes[building_class] == '9b' or bldg_classes[building_class] == '9a':
            target_wall_glazing_U = 2.0
        else:
            if aus_climate_zone[climate_zone] == 2 or aus_climate_zone[climate_zone] == 5:
                if bldg_classes[building_class] == 3 or bldg_classes[building_class] == '9c' or bldg_classes[building_class] == '9a ward':
                    target_wall_glazing_U = 2.0 
            elif aus_climate_zone[climate_zone] == 1 or aus_climate_zone[climate_zone] == 3 or aus_climate_zone[climate_zone] == 4 or aus_climate_zone[climate_zone] == 6 or aus_climate_zone[climate_zone] == 7:
                if bldg_classes[building_class] == 3 or bldg_classes[building_class] == '9c' or bldg_classes[building_class] == '9a ward':
                    target_wall_glazing_U = 1.1
            elif aus_climate_zone[climate_zone] == 8:
                if bldg_classes[building_class] == 3 or bldg_classes[building_class] == '9c' or bldg_classes[building_class] == '9a ward':
                    target_wall_glazing_U = 0.9

        u_value_glazing = []
        wall_glazing_u_value = []
        solar_admittance_single = []
        sum_UA = []

        for i in range(0,4):
            x = (target_wall_glazing_U*model_faces_vertical['Face Area (m2)'].iloc[i]-(Wall_U_Value[i]*(model_faces_vertical['Face Area (m2)'].iloc[i]-model_apertures['Aperture Area (m2)'].iloc[i])))/model_apertures['Aperture Area (m2)'].iloc[i]
            UA = ((1/ex_wall_dts)*model_faces_vertical['ExWall Area (m2)'].iloc[i]) + (glass_u_dts*model_apertures['Aperture Area (m2)'].iloc[i])
            sum_UA.append(((1/ex_wall_dts)*model_faces_vertical['ExWall Area (m2)'].iloc[i]) + (glass_u_dts*model_apertures['Aperture Area (m2)'].iloc[i]))
            wall_glazing_u_value.append(round(UA/model_faces_vertical['Face Area (m2)'].iloc[i],2))
            
            if x == np.inf:
                u_value_glazing.append(0) #Y66
            else:
                u_value_glazing.append(x)
            #----
            if model_apertures['Aperture Area (m2)'].iloc[i] == 0 :
                solar_admittance_single.append(0)
            elif model_apertures['Aperture Area (m2)'].iloc[i] > 0 :
                solar_admittance_single.append(round((shading_multi*glass_shgc_dts*model_apertures['Aperture Area (m2)'].iloc[i]) / model_faces_vertical['Face Area (m2)'].iloc[i],2))

    
        dts_glazing_U = DataFrame([u_value_glazing,model_apertures['Aperture Area (m2)'].values]).transpose()
        dts_glazing_U.rename(columns = {0:'U-Value Glazing', 1:'Vision Area'},index = {0:'East',1:'North',2:'South',3:'West'}, inplace= True)
        
        dts_glazing_U['Area Weighted U-value Glazing'] = dts_glazing_U['U-Value Glazing']*model_apertures['Aperture Area (m2)']
        Reference_Building_glazing_U_value = round(dts_glazing_U['Area Weighted U-value Glazing'].sum()/model_apertures['Aperture Area (m2)'].sum(),2)
        Reference_Building_wall_U_value = (Wall_U_Value[0]*model_faces_vertical['ExWall Area (m2)'].iloc[0]+Wall_U_Value[1]*model_faces_vertical['ExWall Area (m2)'].iloc[1]+Wall_U_Value[2]*model_faces_vertical['ExWall Area (m2)'].iloc[2]+Wall_U_Value[3]*model_faces_vertical['ExWall Area (m2)'].iloc[3]) / model_faces_vertical['ExWall Area (m2)'].sum()

        #Solar admittancce targets
        if bldg_classes[building_class] == 2 or bldg_classes[building_class] == 5 or bldg_classes[building_class] == 6 or bldg_classes[building_class] == 7 or bldg_classes[building_class] == 8 or bldg_classes[building_class] == '9b' or bldg_classes[building_class] == '9a':
            if aus_climate_zone[climate_zone] ==2 or aus_climate_zone[climate_zone] == 4 or aus_climate_zone[climate_zone] == 5 or aus_climate_zone[climate_zone] == 6 or aus_climate_zone[climate_zone] == 7:
                solar_admittance = {'East':0.13, 'North':0.13,'South':0.13, 'West':0.13}
            elif aus_climate_zone[climate_zone] ==1:
                solar_admittance = {'East':0.12, 'North':0.12,'South':0.12, 'West':0.12}
            elif aus_climate_zone[climate_zone] ==3:
                solar_admittance = {'East':0.16, 'North':0.16,'South':0.16, 'West':0.16}
            elif aus_climate_zone[climate_zone] ==8:
                solar_admittance = {'East':0.20, 'North':0.20,'South':0.42, 'West':0.36}
        elif bldg_classes[building_class] == 3 or bldg_classes[building_class] == '9c' or bldg_classes[building_class] == '9a ward':
            if aus_climate_zone[climate_zone] ==1:
                solar_admittance = {'East':0.07, 'North':0.07,'South':0.10, 'West':0.07}
            elif aus_climate_zone[climate_zone] == 3 or aus_climate_zone[climate_zone] == 4 or aus_climate_zone[climate_zone] ==6:
                solar_admittance = {'East':0.07, 'North':0.07,'South':0.07, 'West':0.07}
            elif aus_climate_zone[climate_zone] == 2 or aus_climate_zone[climate_zone] == 5 :
                solar_admittance = {'East':0.10, 'North':0.10,'South':0.10, 'West':0.10}
            elif aus_climate_zone[climate_zone] == 7:
                solar_admittance = {'East':0.07, 'North':0.07,'South':0.08, 'West':0.07}
            elif aus_climate_zone[climate_zone] == 8:
                solar_admittance = {'East':0.08, 'North':0.08,'South':0.08, 'West':0.08}
        if bldg_classes[building_class] == 2 or bldg_classes[building_class] == 5 or bldg_classes[building_class] == 6 or bldg_classes[building_class] == 7 or bldg_classes[building_class] == 8 or bldg_classes[building_class] == '9b' or bldg_classes[building_class] == '9a':
            if aus_climate_zone[climate_zone] == 1:
                solar_admittance_weight_coe = {'East':1.39, 'North':1.47,'South':1, 'West':1.41}
            elif aus_climate_zone[climate_zone] == 2:
                solar_admittance_weight_coe = {'East':1.58, 'North':1.95,'South':1, 'West':1.68}
            elif aus_climate_zone[climate_zone] == 3:
                solar_admittance_weight_coe = {'East':1.63, 'North':1.95,'South':1, 'West':1.65}
            elif aus_climate_zone[climate_zone] == 4:
                solar_admittance_weight_coe = {'East':1.72, 'North':2.05,'South':1, 'West':1.69}
            elif aus_climate_zone[climate_zone] == 5:
                solar_admittance_weight_coe = {'East':1.72, 'North':2.28,'South':1, 'West':1.75}
            elif aus_climate_zone[climate_zone] == 6:
                solar_admittance_weight_coe = {'East':1.62, 'North':2.12,'South':1, 'West':1.67}
            elif aus_climate_zone[climate_zone] == 7:
                solar_admittance_weight_coe = {'East':1.84, 'North':2.4,'South':1, 'West':1.92}
            elif aus_climate_zone[climate_zone] == 8:
                solar_admittance_weight_coe = {'East':1.92, 'North':1.88,'South':1, 'West':1.25}
        elif bldg_classes[building_class] == 3 or bldg_classes[building_class] == '9c' or bldg_classes[building_class] == '9a ward':
            if aus_climate_zone[climate_zone] == 1:
                solar_admittance_weight_coe = {'East':1.3, 'North':1.47,'South':1, 'West':1.37}
            elif aus_climate_zone[climate_zone] == 2:
                solar_admittance_weight_coe = {'East':1.49, 'North':1.77,'South':1, 'West':1.54}
            elif aus_climate_zone[climate_zone] == 3:
                solar_admittance_weight_coe = {'East':1.48, 'North':1.72,'South':1, 'West':1.5}
            elif aus_climate_zone[climate_zone] == 4:
                solar_admittance_weight_coe = {'East':1.37, 'North':1.55,'South':1, 'West':1.36}
            elif aus_climate_zone[climate_zone] == 5:
                solar_admittance_weight_coe = {'East':1.48, 'North':1.88,'South':1, 'West':1.52}
            elif aus_climate_zone[climate_zone] == 6:
                solar_admittance_weight_coe = {'East':1.28, 'North':1.52,'South':1, 'West':1.33}
            elif aus_climate_zone[climate_zone] == 7:
                solar_admittance_weight_coe = {'East':1.35, 'North':1.6,'South':1, 'West':1.4}
            elif aus_climate_zone[climate_zone] == 8:
                solar_admittance_weight_coe = {'East':1.26, 'North':1.24,'South':1, 'West':1.05}
            
        
        
        dts_shgc_single = {'East':[round(solar_admittance['East']/(shading_multi*(model_faces_vertical['WWR (%)'].iloc[0]/100)),2)],
                    'North':[round(solar_admittance['North']/(shading_multi*(model_faces_vertical['WWR (%)'].iloc[1]/100)),2)],
                    'South':[round(solar_admittance['South']/(shading_multi*(model_faces_vertical['WWR (%)'].iloc[2]/100)),2)],
                    'West':[round(solar_admittance['West']/(shading_multi*(model_faces_vertical['WWR (%)'].iloc[3]/100)),2)]}
        
        
        dts_shgc_single = DataFrame.from_dict(dts_shgc_single).transpose()
        dts_shgc_single.rename(columns = {0:'SHGC'},inplace= True)
        

        #SA COE based on WWR%
        if model_faces_vertical['WWR (%)'].iloc[0] < 20: #East
            solar_admittance_weight_coe_east = 0
        elif model_faces_vertical['WWR (%)'].iloc[0] >= 20:
            solar_admittance_weight_coe_east = solar_admittance_weight_coe['East']
        if model_faces_vertical['WWR (%)'].iloc[1] < 20: #North
            solar_admittance_weight_coe_north = 0
        elif model_faces_vertical['WWR (%)'].iloc[1] >= 20:
            solar_admittance_weight_coe_north = solar_admittance_weight_coe['North']
        if model_faces_vertical['WWR (%)'].iloc[2] < 20: #South
            solar_admittance_weight_coe_south = 0
        elif model_faces_vertical['WWR (%)'].iloc[2] >= 20:
            solar_admittance_weight_coe_south = solar_admittance_weight_coe['South']
        if model_faces_vertical['WWR (%)'].iloc[3] < 20: #West
            solar_admittance_weight_coe_west = 0
        elif model_faces_vertical['WWR (%)'].iloc[3] >= 20:
            solar_admittance_weight_coe_west = solar_admittance_weight_coe['West']
        
        #SA based on aperture existance
        if model_faces_vertical['WWR (%)'].iloc[0] == 0: #East
            solar_admittance_weight_coe_east = 0
        elif model_faces_vertical['WWR (%)'].iloc[0] >= 20:
            solar_admittance_weight_coe_east = solar_admittance_weight_coe['East']
        if model_faces_vertical['WWR (%)'].iloc[1] < 20: #North
            solar_admittance_weight_coe_north = 0
        elif model_faces_vertical['WWR (%)'].iloc[1] >= 20:
            solar_admittance_weight_coe_north = solar_admittance_weight_coe['North']
        if model_faces_vertical['WWR (%)'].iloc[2] < 20: #South
            solar_admittance_weight_coe_south = 0
        elif model_faces_vertical['WWR (%)'].iloc[2] >= 20:
            solar_admittance_weight_coe_south = solar_admittance_weight_coe['South']
        if model_faces_vertical['WWR (%)'].iloc[3] < 20: #West
            solar_admittance_weight_coe_west = 0
        elif model_faces_vertical['WWR (%)'].iloc[3] >= 20:
            solar_admittance_weight_coe_west = solar_admittance_weight_coe['West']



        #Total Values
        reference_ac_energy = model_faces_vertical['Face Area (m2)'].iloc[0]*solar_admittance_weight_coe_east*solar_admittance['East'] + model_faces_vertical['Face Area (m2)'].iloc[1]*solar_admittance_weight_coe_north*solar_admittance['North'] + model_faces_vertical['Face Area (m2)'].iloc[2]*solar_admittance_weight_coe_south*solar_admittance['South'] + model_faces_vertical['Face Area (m2)'].iloc[3]*solar_admittance_weight_coe_west*solar_admittance['West']
        dts_shgc_total = reference_ac_energy/((solar_admittance_weight_coe_east*dts_glazing_U['Vision Area'].iloc[0]*shading_multi)+(solar_admittance_weight_coe_north*dts_glazing_U['Vision Area'].iloc[1]*shading_multi)+(solar_admittance_weight_coe_south*dts_glazing_U['Vision Area'].iloc[2]*shading_multi)+(solar_admittance_weight_coe_west*dts_glazing_U['Vision Area'].iloc[3]*shading_multi))
        wall_glazing_value_total = sum(sum_UA) / model_faces_vertical['Face Area (m2)'].sum()
        proposed_ac_energy = model_faces_vertical['Face Area (m2)'].iloc[0]*solar_admittance_weight_coe_east*solar_admittance_single[0]+model_faces_vertical['Face Area (m2)'].iloc[1]*solar_admittance_weight_coe_north*solar_admittance_single[1]+model_faces_vertical['Face Area (m2)'].iloc[2]*solar_admittance_weight_coe_south*solar_admittance_single[2]+model_faces_vertical['Face Area (m2)'].iloc[3]*solar_admittance_weight_coe_west*solar_admittance_single[3]

        st.subheader("Method 1:")
        cols = st.columns(4)
        with cols[0]:
            st.metric("East Wall U-Value(W/m².K", round(Wall_U_Value[0],2))
        with cols[1]:
            st.metric("North Wall U-Value(W/m².K)", round(Wall_U_Value[1],2))
        with cols[2]:
            st.metric("South Wall U-Value(W/m².K)", round(Wall_U_Value[2],2))
        with cols[3]:
            st.metric("West Wall U-Value(W/m².K)", round(Wall_U_Value[3],2))
        cols = st.columns(4)
        with cols[0]:
            if dts_glazing_U['U-Value Glazing'].iloc[0] == 0:
                st.metric("East Glazing U-Value(W/m².K)",None)
            else:
                if dts_glazing_U['U-Value Glazing'].iloc[0] > 5.8:
                    st.metric("East Glazing U-Value(W/m².K)",5.8)
                elif Reference_Building_glazing_U_value < 1.5:
                    st.metric("East Glazing U-Value(W/m².K)",1.5)
                else:
                    st.metric("East Glazing U-Value(W/m².K)", round(dts_glazing_U['U-Value Glazing'].iloc[0],2))
        with cols[1]:
            if dts_glazing_U['U-Value Glazing'].iloc[0] == 0:
                st.metric("North Glazing U-Value(W/m².K)",None)
            else:
                if dts_glazing_U['U-Value Glazing'].iloc[1] > 5.8:
                    st.metric("North Glazing U-Value(W/m².K)",5.8)
                elif Reference_Building_glazing_U_value < 1.5:
                    st.metric("North Glazing U-Value(W/m².K)",1.5)
                else:
                    st.metric("North Glazing U-Value(W/m².K)", round(dts_glazing_U['U-Value Glazing'].iloc[1],2))
        with cols[2]:
            if dts_glazing_U['U-Value Glazing'].iloc[0] == 0:
                st.metric("South Glazing U-Value(W/m².K)",None)
            else:
                if dts_glazing_U['U-Value Glazing'].iloc[2] >= 5.8:
                    st.metric("South Glazing U-Value(W/m².K)",5.8)
                elif Reference_Building_glazing_U_value <= 1.5:
                    st.metric("South Glazing U-Value(W/m².K)",1.5)
                else:
                    st.metric("South Glazing U-Value(W/m².K)", round(dts_glazing_U['U-Value Glazing'].iloc[2],2))
        with cols[3]:
            if dts_glazing_U['U-Value Glazing'].iloc[0] == 0:
                st.metric("West Glazing U-Value(W/m².K)",None)
            else:
                if dts_glazing_U['U-Value Glazing'].iloc[3] > 5.8:
                    st.metric("West Glazing U-Value(W/m².K)",5.8)
                elif Reference_Building_glazing_U_value < 1.5:
                    st.metric("West Glazing U-Value(W/m².K)",1.5)
                else:
                    st.metric("West Glazing U-Value(W/m².K)", round(dts_glazing_U['U-Value Glazing'].iloc[3],2))
        cols = st.columns(4)
        with cols[0]:
            if dts_shgc_single['SHGC'].iloc[0] > 0.81:
                st.metric("East Glazing SHGC",0.81)
            elif dts_shgc_single['SHGC'].iloc[0] == np.inf:
                st.metric("East Glazing SHGC",0)
            elif dts_shgc_single['SHGC'].iloc[0] < 0.16:
                st.metric("East Glazing SHGC",0.16)
            else:
                st.metric("East Glazing SHGC",dts_shgc_single.iloc[0])
        with cols[1]:
            if dts_shgc_single['SHGC'].iloc[1] > 0.81:
                st.metric("North Glazing SHGC",0.81)
            elif dts_shgc_single['SHGC'].iloc[1] == 0:
                st.metric("North Glazing SHGC",0)
            elif dts_shgc_single['SHGC'].iloc[1] < 0.16:
                st.metric("North Glazing SHGC",0.16)
            else: 
                st.metric("North Glazing SHGC",dts_shgc_single.iloc[1])
        with cols[2]:
            if dts_shgc_single['SHGC'].iloc[2] > 0.81:
                st.metric("South Glazing SHGC",0.81)
            elif dts_shgc_single['SHGC'].iloc[2] == np.inf:
                st.metric("South Glazing SHGC",0)
            elif dts_shgc_single['SHGC'].iloc[2] < 0.16:
                st.metric("South Glazing SHGC",0.16)
            else: 
                st.metric("South Glazing SHGC",dts_shgc_single.iloc[2])
        with cols[3]:
            if dts_shgc_single['SHGC'].iloc[3] > 0.81:
                st.metric("West Glazing SHGC",0.81)
            elif dts_shgc_single['SHGC'].iloc[3] == np.inf:
                st.metric("West Glazing SHGC",0)
            elif dts_shgc_single['SHGC'].iloc[3] < 0.16:
                st.metric("West Glazing SHGC",0.16)
            else: 
                st.metric("West Glazing SHGC",dts_shgc_single.iloc[3])

        def CheckForLess(list, val):
 
                # traverse in the list
                for x in list:

                    # compare with all the
                    # values with value
                    if val <= x:
                        return False
                return True

        cols = st.columns([5,1,5])

        with cols[0]:
            #Bar chart Wall glazing U value
            wall_glazing_u_bar = go.Figure(data=[go.Bar(x=['East','North', 'South','West'],
                                                    y=wall_glazing_u_value,marker_color='lightslategray',text=wall_glazing_u_value)])
        
            wall_glazing_u_bar.update_layout(
                        yaxis = dict(title = "Wall Glazing U-Value W/m².K"),
                        shapes=[
                            {
                                'type': 'line',
                                'xref': 'paper',
                                'x0': 0,
                                'y0': target_wall_glazing_U,
                                'x1': 1,
                                'y1': target_wall_glazing_U,
                                'line': {
                                    'color': 'rgb(50, 171, 96)',
                                    'width': 2,
                                    'dash': 'dash',
                                },
                            },
                        ],
                            )
            wall_glazing_u_bar.add_annotation(
                x=0
                , y=target_wall_glazing_U
                , text=f'DtS Threshold'
                , yanchor='bottom'
                , showarrow=True
                , arrowhead=1
                , arrowsize=1
                , arrowwidth=2
                , arrowcolor="#636363"
                , ax=-20
                , ay=-30
                , font=dict(size=15, color="black", family="Arial")
                , align="left"
                ,)
            
            st.plotly_chart(wall_glazing_u_bar, use_container_width=True)
            wall_glazing_u_bar.write_image("wall_glazing_u_bar.png")


            if (CheckForLess(wall_glazing_u_value, target_wall_glazing_U)):
                st.subheader("**:green[Compliant Solution]**")
                method1_wall_glazing = 'Compliant Solution'
            else:
                st.subheader("**:red[Non-Compliant Solution]**")
                method1_wall_glazing = 'Non-Compliant Solution'
                st.markdown("**Recommendation: Increasing  external wall r-value OR reducing glass U-value**")

        with cols[1]:
            ""

        with cols[2]:
            #Bar chart solar admittance
            sa_bar = go.Figure(data=[go.Bar(x=list(solar_admittance.keys()),
                                                        y=solar_admittance_single,marker_color='lightslategray',text=solar_admittance_single)])
            
            sa_bar.update_layout(
                        yaxis = dict(title = "Solar Admittance"), showlegend = False)
            
            sa_bar.add_trace(
                go.Scatter(
                    x=list(solar_admittance.keys()),
                    y=list(solar_admittance.values()),
                    name="DtS Threshold",
                    mode='lines+markers',   
                ),
            )

            sa_bar.add_annotation(
                x=0
                , y=solar_admittance['North']
                , text=f'DtS Threshold'
                , yanchor='bottom'
                , showarrow=True
                , arrowhead=1
                , arrowsize=1
                , arrowwidth=2
                , arrowcolor="#636363"
                , ax=-20
                , ay=-30
                , font=dict(size=15, color="black", family="Arial")
                , align="right"
                ,)
            
            st.plotly_chart(sa_bar, use_container_width=True)
            sa_bar.write_image("sa_bar.png")

            if solar_admittance_single <= list(solar_admittance.values()):
                st.subheader("**:green[Compliant Solution]**")
                method1_sa = 'Compliant Solution'
            else:
                st.subheader("**:red[Non-Compliant Solution]**")
                method1_sa = 'Non-Compliant Solution'
                st.markdown("**Recommendation: Reducing glass SHGC**")

        st.subheader("Method 2:")
        cols = st.columns(3)
        with cols[0]:
            st.metric("Reference Building Wall U-value",round(Reference_Building_wall_U_value,2))
        with cols[1]:  
            if Reference_Building_glazing_U_value > 5.8:
                st.metric("Reference Building Glazing U-value",5.8)
            elif Reference_Building_glazing_U_value < 1.5:
                st.metric("Reference Building Glazing U-value",1.5)
            else:
                st.metric("Reference Building Glazing U-value",Reference_Building_glazing_U_value)
        with cols[2]:
            if dts_shgc_total > 0.81:
                st.metric("Reference Building Glazing SHGC",0.81)
            elif dts_shgc_total == np.inf:
                st.metric("Reference Building Glazing SHGC",0)
            elif dts_shgc_total < 0.16:
                st.metric("Reference Building Glazing SHGC",0.16)
            else:
                st.metric("Reference Building Glazing SHGC",round(dts_shgc_total,2))

        cols = st.columns([3,3,3])

        with cols[0]:
            #Bar chart Wall glazing U value total
            wall_glazing_u_total_bar = go.Figure(data=[go.Bar(x=['Proposed Design','DtS Reference'],
                                                    y=[round(wall_glazing_value_total,2),target_wall_glazing_U],marker_color=['lightgreen','lightslategray'],text=[round(wall_glazing_value_total,2),target_wall_glazing_U])])
            wall_glazing_u_total_bar.update_layout(
                        yaxis = dict(title = "Wall Glazing U-Value W/m².K Total"))

            st.plotly_chart(wall_glazing_u_total_bar, use_container_width=True)
            wall_glazing_u_total_bar.write_image("wall_glazing_u_total_bar.png")

            if wall_glazing_value_total <= target_wall_glazing_U:
                st.subheader("**:green[Compliant Solution]**")
                method2_wall_glazing = 'Compliant Solution'
            else:
                st.subheader("**:red[Non-Compliant Solution]**")
                method2_wall_glazing = 'Non-Compliant Solution'
                st.markdown("**Recommendation: Increasing  external wall r-value OR reducing glass U-value**")

        with cols[1]:
            ""

        with cols[2]:
            #Bar chart AC energy total
            AC_energy = go.Figure(data=[go.Bar(x=['Proposed Design','DtS Reference'],
                                                    y=[round(proposed_ac_energy,2),round(reference_ac_energy,2)],marker_color=['lightgreen','lightslategray'],text=[round(proposed_ac_energy,2),round(reference_ac_energy,2)])])
            AC_energy.update_layout(
                        yaxis = dict(title = "AC Energy Value"))

            st.plotly_chart(AC_energy, use_container_width=True)
            AC_energy.write_image("AC_energy.png")

            if proposed_ac_energy <= reference_ac_energy:
                st.subheader("**:green[Compliant Solution]**")
                method2_ac_energy = 'Compliant Solution'
            else:
                st.subheader("**:red[Non-Compliant Solution]**")
                method2_ac_energy  = 'Non-Compliant Solution'
                st.markdown("**Recommendation: Reducing glass SHGC**")
        

        with cols[2]:
            ""
    #Generating the report
    if method1_wall_glazing == 'Compliant Solution' and method1_sa == 'Compliant Solution':
        method1_compliance = 'Compliant Solution'
    else:
        method1_compliance = 'Non-Compliant Solution'

    if method2_wall_glazing == 'Compliant Solution' and method2_ac_energy == 'Compliant Solution':
        method2_compliance = 'Compliant Solution'
    else:
        method2_compliance = 'Non-Compliant Solution'

    NCC19 = Document()
    section = NCC19.sections[0]
    section.page_height = Mm(350)
    section.page_width = Mm(297)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    style = NCC19.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    x = 2.5
    h_res = x
    w_res = 2*x

    cols = st.columns(4)
    with cols[0]:
        project_name = st.text_input("**Building Name / Address:**", value = f'{callback_once().display_name}')
    with cols[1]:
        name = st.text_input("**Your Name / Position:**", value = 'Your Name/Position is required!')
    with cols[2]:
        today = datetime.datetime.now()
        Date = st.date_input("**Date:**", value = today)
    with cols[3]:
        levels = st.text_input("**Storeys Above Ground:**", value = 1)

    header = NCC19.sections[0].header

    header.paragraphs[0].text = f'NCC19 Facade Calculator | Date: {Date} | Approved by: {name}'
    NCC19.add_heading(f'NCC 2019 DtS Project Summary for {project_name}', 0)
    NCC19.add_paragraph('The summary below provides an overview of where compliance has been achieved for Specification J1.5a - Calculation of U-Value and solar admittance - Method 1 (Single Aspect) and Method 2 (Multiple Apects).')
    NCC19.add_paragraph(f'This {building_class} located in {building_state} with a {climate_zone} in {levels} level(s)')
    
    paragraph = NCC19.add_paragraph()
    entire_model = pd.concat([round(model_faces_vertical,2),round(model_apertures,2)], axis = 1)
    entire_model.reset_index(inplace = True)
    entire_model.rename(columns={'index':'Face Direction'}, inplace= True)
    t = NCC19.add_table(entire_model.shape[0]+1, entire_model.shape[1])

    # Adding style to a table 
    t.style = 'Medium Shading 1'

    # add the header rows.
    for j in range(entire_model.shape[-1]):
        t.cell(0,j).text = entire_model.columns[j]

    # add the rest of the data frame
    for i in range(entire_model.shape[0]):
        for j in range(entire_model.shape[-1]):
            t.cell(i+1,j).text = str(entire_model.values[i,j])


    NCC19.add_heading('METHOD 1:', 1)
    paragraph = NCC19.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('wall_glazing_u_bar.png', width=Inches(w_res), height= Inches(h_res))
    run_2 = paragraph.add_run()
    run_2.add_picture('sa_bar.png', width=Inches(w_res), height= Inches(h_res))
    
    NCC19.add_heading('METHOD 2:', 1)
    paragraph = NCC19.add_paragraph()
    run = paragraph.add_run()
    run.add_picture('wall_glazing_u_total_bar.png', width=Inches(w_res), height= Inches(h_res))
    run_2 = paragraph.add_run()
    run_2.add_picture('AC_energy.png', width=Inches(w_res), height= Inches(h_res))


    p0 = NCC19.add_paragraph(f'The proposed thermal envelope performance includes an External Wall R-value of ')
    
    p0.add_run(f'{ex_wall_dts}, and Glass U-value & SHGC of {glass_u_dts} / {glass_shgc_dts}').bold= True
    p0.add_run(' which indicates the proposed design as ')
    p0.add_run(f'{method1_compliance}').bold= True
    p0.add_run(' against NCC19 DtS Reference Method 1, and ')
    p0.add_run(f'{method2_compliance}').bold= True
    p0.add_run(' against NCC19 DtS Reference Method 2.')

    # NCC19.add_heading(f'Project Details', 1)

    buffer = io.BytesIO()
    NCC19.save(buffer)
    export_as_word = st.download_button(
            label="Generate NCC19 Facade Calculator Report.docx",
            data=buffer,
            file_name=f'NCC19 Glass Facade Calculator - {project_name}.docx',
            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
